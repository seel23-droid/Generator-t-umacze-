import os
import tempfile
import streamlit as st
from docx import Document
from pdf2docx import Converter
from google import genai
from google.genai import types

# ---------------------------------------------------------------------------
# STRONA I KONFIGURACJA STREAMLIT
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Generator Tłumaczeń Dokumentów",
    page_icon="📄",
    layout="centered"
)

st.title("📄 Generator Tłumaczeń Dokumentów")
st.markdown("""
Aplikacja tłumaczy pliki **DOCX** oraz **PDF** z języka angielskiego na język polski, 
starając się zachować **oryginalny układ stron, tabele, formatowanie oraz obrazy**.
""")

# Pasek boczny - Klucz API
st.sidebar.header("⚙️ Konfiguracja")
api_key_input = st.sidebar.text_input(
    "Wprowadź Gemini API Key:",
    type="password",
    help="Pobierz klucz z Google AI Studio (https://aistudio.google.com/)"
)

# Sprawdzenie czy klucz jest dostępny (z paska lub zmiennej środowiskowej)
api_key = api_key_input or os.environ.get("GEMINI_API_KEY")

# ---------------------------------------------------------------------------
# INSTUKCJA DLA SYSTEMU
# ---------------------------------------------------------------------------
SYSTEM_INSTRUCTION = """
Jesteś profesjonalnym tłumaczem dokumentów z języka angielskiego na język polski.
Twoim zadaniem jest przetłumaczenie podanego tekstu na naturalny, poprawny język polski.
Zasady:
1. Zachowaj oryginalne znaczenie i kontekst techniczny/biznesowy.
2. Zachowaj wszelkie znaki specjalne, formatowanie liczb oraz kody/identyfikatory.
3. Zwracaj WYŁĄCZNIE przetłumaczony tekst — nie dodawaj żadnych komentarzy, wstępów ani wyjaśnień.
"""

def get_gemini_client(key: str):
    return genai.Client(api_key=key)

def translate_text(client: genai.Client, text: str) -> str:
    """Tłumaczy fragment tekstu przy użyciu Gemini API."""
    if not text or not text.strip():
        return text

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=text,
            config=types.GenerateContentConfig(
                system_instruction=SYSTEM_INSTRUCTION,
                temperature=0.2,
            ),
        )
        return response.text.strip() if response.text else text
    except Exception as e:
        st.warning(f"Błąd tłumaczenia fragmentu tekstu: {e}")
        return text

# ---------------------------------------------------------------------------
# TŁUMACZENIE PLIKÓW DOCX
# ---------------------------------------------------------------------------
def translate_docx_file(client: genai.Client, input_path: str, output_path: str, progress_bar, status_text):
    """Tłumaczy plik DOCX modyfikując jego treść in-place."""
    doc = Document(input_path)

    # Policz łączną liczbę elementów do przetłumaczenia
    paragraphs = doc.paragraphs
    tables = doc.tables

    total_items = len(paragraphs)
    for t in tables:
        for r in t.rows:
            total_items += len(r.cells)

    processed_items = 0

    # 1. Akapity główne
    status_text.text("Tłumaczenie akapitów głównych...")
    for p in paragraphs:
        if p.text.strip():
            translated = translate_text(client, p.text)
            if p.runs:
                p.runs[0].text = translated
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = translated
        
        processed_items += 1
        if total_items > 0:
            progress_bar.progress(min(processed_items / total_items, 1.0))

    # 2. Tabele
    status_text.text("Tłumaczenie komórek w tabelach...")
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        translated = translate_text(client, p.text)
                        if p.runs:
                            p.runs[0].text = translated
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.text = translated
                
                processed_items += 1
                if total_items > 0:
                    progress_bar.progress(min(processed_items / total_items, 1.0))

    doc.save(output_path)

# ---------------------------------------------------------------------------
# INTERFEJS I OBSŁUGA PLIKÓW
# ---------------------------------------------------------------------------
uploaded_file = st.file_uploader(
    "Wgraj dokument (PDF lub DOCX):", 
    type=["pdf", "docx"],
    help="Obsługiwane są pliki w formacie PDF oraz Microsoft Word (.docx)"
)

if uploaded_file is not None:
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    file_name_without_ext = os.path.splitext(uploaded_file.name)[0]
    
    st.info(f"Wgranym plik: **{uploaded_file.name}** ({file_ext.upper()})")

    if st.button("🚀 Rozpocznij tłumaczenie", type="primary"):
        if not api_key:
            st.error("Proszę wprowadzić klucz API Gemini w panelu bocznym!")
        else:
            try:
                client = get_gemini_client(api_key)
                progress_bar = st.progress(0.0)
                status_text = st.empty()

                with tempfile.TemporaryDirectory() as temp_dir:
                    input_file_path = os.path.join(temp_dir, uploaded_file.name)
                    output_docx_path = os.path.join(temp_dir, f"{file_name_without_ext}_PL.docx")

                    # Zapisujemy wgrany plik w katalogu tymczasowym
                    with open(input_file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())

                    # Proces dla plików PDF
                    if file_ext == ".pdf":
                        status_text.text("Krok 1/2: Konwersja PDF do formatu DOCX (odtwarzanie układu)...")
                        temp_docx_from_pdf = os.path.join(temp_dir, "temp_conv.docx")
                        
                        cv = Converter(input_file_path)
                        cv.convert(temp_docx_from_pdf, start=0, end=None)
                        cv.close()

                        status_text.text("Krok 2/2: Tłumaczenie treści...")
                        translate_docx_file(client, temp_docx_from_pdf, output_docx_path, progress_bar, status_text)

                    # Proces dla plików DOCX
                    elif file_ext == ".docx":
                        status_text.text("Tłumaczenie treści dokumentu Word...")
                        translate_docx_file(client, input_file_path, output_docx_path, progress_bar, status_text)

                    status_text.text("Gotowe!")
                    progress_bar.progress(1.0)
                    st.success("Tłumaczenie zakończone sukcesem!")

                    # Odczyt wygenerowanego pliku do przycisku pobierania
                    with open(output_docx_path, "rb") as translated_file:
                        st.download_button(
                            label="📥 Pobierz przetłumaczony dokument (.docx)",
                            data=translated_file.read(),
                            file_name=f"{file_name_without_ext}_PL.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(f"Wystąpił błąd podczas przetwarzania pliku: {e}")
